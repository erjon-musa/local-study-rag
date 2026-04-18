"""
Document loaders for the RAG ingestion pipeline.

Supports: PDF, Markdown, TXT, DOCX, HTML
Each loader returns a list of Document objects with text and metadata.

For scanned PDF pages (image-only, no extractable text):
  1. Primary: LightOnOCR-2-1B running locally on Apple Silicon (MPS)
     — layout-aware, no network round-trip, much higher quality.
  2. Fallback: Gemma 4 multimodal OCR via LM Studio over Tailscale.

Toggle the primary with USE_LOCAL_OCR=true|false in .env.
"""
from __future__ import annotations

import base64
import io
import math
import os
import time
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Tuple

import openai

LMSTUDIO_BASE_URL = os.getenv("LMSTUDIO_BASE_URL", "http://localhost:1234/v1")
LMSTUDIO_MODEL = os.getenv("LMSTUDIO_MODEL", "google_gemma-4-26b-a4b-it")

# Local OCR (LightOnOCR-2-1B on MPS)
USE_LOCAL_OCR = os.getenv("USE_LOCAL_OCR", "true").lower() in ("1", "true", "yes")
LOCAL_OCR_MODEL_NAME = os.getenv("LOCAL_OCR_MODEL", "lightonai/LightOnOCR-2-1B")
OCR_MAX_IMAGE_DIM = int(os.getenv("OCR_MAX_IMAGE_DIM", "1024"))

# Lazy-initialized local OCR model — populated on first use to avoid loading
# 2 GB of weights into every Python process that imports this module.
_OCR_MODEL = None
_OCR_PROCESSOR = None
_OCR_DEVICE = None
_OCR_DTYPE = None

# OCR prompt for Gemma4 multimodal — tuned for academic documents
OCR_PROMPT = (
    "Extract ALL text from this document page exactly as written. "
    "Include headings, questions, formulas, and any mathematical notation. "
    "Preserve the structure (numbered questions, sub-parts like (i), (ii), etc.). "
    "Output only the extracted text, nothing else. "
    "Do not use internal reasoning. Respond directly."
)


@dataclass
class Document:
    """A loaded document with text content and metadata."""
    text: str
    metadata: dict = field(default_factory=dict)


def _ensure_local_ocr_model():
    """Lazy-load LightOnOCR model on first use. Idempotent."""
    global _OCR_MODEL, _OCR_PROCESSOR, _OCR_DEVICE, _OCR_DTYPE
    if _OCR_MODEL is not None:
        return

    import torch
    from transformers import LightOnOcrForConditionalGeneration, LightOnOcrProcessor

    _OCR_DEVICE = "mps" if torch.backends.mps.is_available() else "cpu"
    # MPS has spotty bf16 support; fp32 is the documented choice on Apple Silicon
    _OCR_DTYPE = torch.float32 if _OCR_DEVICE == "mps" else torch.bfloat16

    print(f"    Loading local OCR model {LOCAL_OCR_MODEL_NAME} on {_OCR_DEVICE} ({_OCR_DTYPE})...")
    _OCR_PROCESSOR = LightOnOcrProcessor.from_pretrained(LOCAL_OCR_MODEL_NAME)
    _OCR_MODEL = LightOnOcrForConditionalGeneration.from_pretrained(
        LOCAL_OCR_MODEL_NAME, torch_dtype=_OCR_DTYPE
    ).to(_OCR_DEVICE)
    print(f"    Local OCR ready.")


def _ocr_page_local(page_image_bytes: bytes) -> str:
    """
    OCR a page image using LightOnOCR-2-1B on the local MPS device.

    Resizes large images to keep Pixtral's O(n²) attention from OOMing on
    16 GB unified memory. Returns extracted text or empty string on failure.
    """
    try:
        import torch
        from PIL import Image

        _ensure_local_ocr_model()

        pil_image = Image.open(io.BytesIO(page_image_bytes))
        if max(pil_image.size) > OCR_MAX_IMAGE_DIM:
            scale = OCR_MAX_IMAGE_DIM / max(pil_image.size)
            new_size = (int(pil_image.size[0] * scale), int(pil_image.size[1] * scale))
            pil_image = pil_image.resize(new_size, Image.LANCZOS)

        conversation = [{"role": "user", "content": [{"type": "image", "image": pil_image}]}]
        inputs = _OCR_PROCESSOR.apply_chat_template(
            conversation,
            add_generation_prompt=True,
            tokenize=True,
            return_dict=True,
            return_tensors="pt",
        )
        inputs = {
            k: v.to(device=_OCR_DEVICE, dtype=_OCR_DTYPE) if v.is_floating_point() else v.to(_OCR_DEVICE)
            for k, v in inputs.items()
        }

        if _OCR_DEVICE == "mps":
            torch.mps.empty_cache()

        output_ids = _OCR_MODEL.generate(**inputs, max_new_tokens=2048)
        generated_ids = output_ids[0, inputs["input_ids"].shape[1]:]
        return _OCR_PROCESSOR.decode(generated_ids, skip_special_tokens=True).strip()
    except Exception as e:
        print(f"    ⚠ Local OCR failed: {e}")
        return ""


def _ocr_page_with_gemma(page_image_bytes: bytes) -> str:
    """
    Fallback OCR via Gemma 4 multimodal on LM Studio.

    Used when USE_LOCAL_OCR=false, or when local OCR returns empty.
    """
    b64_image = base64.b64encode(page_image_bytes).decode("utf-8")

    try:
        client = openai.OpenAI(
            base_url=LMSTUDIO_BASE_URL,
            api_key="lmstudio-link",
        )

        response = client.chat.completions.create(
            model=LMSTUDIO_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": OCR_PROMPT},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{b64_image}"
                            },
                        },
                    ],
                }
            ],
            temperature=0.1,
            max_tokens=4096,
            extra_body={"thinking": {"type": "disabled"}},
        )

        message = response.choices[0].message
        text = message.content or ""
        if not text:
            reasoning = getattr(message, "reasoning_content", None) or ""
            text = reasoning

        return text.strip()
    except Exception as e:
        print(f"    ⚠ Gemma OCR failed: {e}")
        return ""


def _ocr_page(page_image_bytes: bytes) -> Tuple[str, str]:
    """
    Dispatch OCR to local model first, fall back to Gemma if local fails.

    Returns (text, method_used) where method is 'ocr_local_lighton' or 'ocr_gemma4'.
    """
    if USE_LOCAL_OCR:
        text = _ocr_page_local(page_image_bytes)
        if text:
            return text, "ocr_local_lighton"
        print(f"    ⤷ Falling back to Gemma OCR over LM Studio...")

    text = _ocr_page_with_gemma(page_image_bytes)
    return text, "ocr_gemma4"


def _pixmap_grayscale_stddev(page) -> float:
    """
    Render `page` at dpi=72 and return the grayscale stddev of its pixels
    on a 0-255 scale.

    Why stddev? A genuinely blank page (all-white pixmap) has stddev ≈ 0.
    A page with *any* visual content — printed text rendered as pixels,
    embedded diagrams, scanned ink — has meaningful variance. We compute
    this without numpy: a single pass through the pixmap bytes accumulates
    sum and sum-of-squares, from which stddev falls out in O(n) with
    constant memory.

    Returns 0.0 on any failure so the caller treats such pages as blank
    (safe default — the original `page.get_images()` path still fires
    if there are registered images).
    """
    try:
        # dpi=72 is one device-pixel-per-PDF-point — cheapest faithful render.
        # colorspace=GRAY means one byte per pixel; no RGB->gray conversion needed.
        import fitz  # noqa: F401  # ensure PyMuPDF is importable before touching `page`

        pixmap = page.get_pixmap(dpi=72, colorspace="gray", alpha=False)
    except Exception:
        return 0.0

    try:
        n = pixmap.width * pixmap.height
        if n <= 0:
            return 0.0

        samples = pixmap.samples  # bytes: one byte per pixel at colorspace=gray
        # If PyMuPDF didn't honor GRAY (very old versions), fall back to raw bytes/length.
        if len(samples) < n:
            return 0.0

        # Stride-sample every ~stride-th pixel. At dpi=72 an A4 page has
        # ~500k grayscale bytes; iterating all of them in pure Python costs
        # ~50ms per page. A stride that yields ~4096 samples is statistically
        # plenty to distinguish stddev≈0 (blank) from stddev>10 (any content)
        # and costs <1ms.
        target_samples = 4096
        stride = max(1, n // target_samples)

        total = 0
        total_sq = 0
        count = 0
        for i in range(0, n, stride):
            v = samples[i]
            total += v
            total_sq += v * v
            count += 1

        if count == 0:
            return 0.0
        mean = total / count
        variance = max(0.0, (total_sq / count) - (mean * mean))
        return math.sqrt(variance)
    except Exception:
        return 0.0


def _page_has_visual_content(page) -> bool:
    """
    Return True when a PDF page is worth handing to OCR.

    Two paths cover the real-world cases we've seen:
      (a) `page.get_images()` non-empty — the page has registered raster
          images (the PyMuPDF-native "there's an image here" signal).
      (b) The page renders to a pixmap whose grayscale stddev exceeds 10
          on the 0-255 scale — catches scans stored as CropBox-clipped
          pixmaps that aren't registered images (the ELEC 472 lectures).

    Returns False for genuinely blank white pages (stddev ≈ 0-2), so we
    don't waste OCR cycles on title-card / spacer pages.
    """
    try:
        if page.get_images():
            return True
    except Exception:
        # Defensive: some PDFs have malformed xref tables. Fall through to stddev.
        pass

    # Threshold 10 chosen empirically: blank pixmaps measure 0-2; pages with
    # any printed text or diagrams measure 30+. 10 is a comfortable margin.
    return _pixmap_grayscale_stddev(page) > 10.0


@dataclass
class PageLoadResult:
    """
    Per-page outcome of PDF loading. Surfaced up to the pipeline so we can
    report OCR successes, OCR failures, and skipped blank pages per file
    instead of silently dropping them.
    """
    page: int
    extraction_method: str  # "text" | "ocr_local_lighton" | "ocr_gemma4" | "skipped_blank" | "failed"
    text_length: int
    duration_s: float = 0.0
    error: str = ""


@dataclass
class PdfLoadStats:
    """Aggregate stats for a single PDF load — attached to pipeline results."""
    ocr_pages: int = 0
    ocr_failed_pages: int = 0
    skipped_blank_pages: int = 0
    pages: List[PageLoadResult] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


def load_pdf(
    path: Path,
    course: str = "",
    category: str = "",
    return_stats: bool = False,
):
    """
    Load a PDF file, returning one Document per page.

    For pages with no extractable text but meaningful visual content
    (`_page_has_visual_content`), falls back to OCR via the LightOnOCR →
    Gemma chain. Blank pages are recorded as `skipped_blank` and OCR
    failures are surfaced with their exception class name.

    Parameters
    ----------
    return_stats : bool
        When True, return `(docs, PdfLoadStats)`. When False (default),
        return just `List[Document]` — preserves the existing call sites
        in pipeline/tests that don't know about the stats.
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError("pymupdf is required: pip install pymupdf")

    docs: List[Document] = []
    stats = PdfLoadStats()
    pdf = fitz.open(str(path))
    ocr_count = 0

    for page_num in range(len(pdf)):
        page = pdf[page_num]
        text = page.get_text("text").strip()
        extraction_method = "text"
        page_error = ""

        # Broadened OCR trigger: fires on any page with visual content (images
        # OR non-blank pixmap) — catches scanned PDFs whose pages are stored
        # as clipped pixmaps without registered images. `DISABLE_OCR` is still
        # a hard kill switch for dev iterations where we don't want MPS churn.
        should_ocr = (
            not text
            and "DISABLE_OCR" not in os.environ
            and _page_has_visual_content(page)
        )

        page_duration = 0.0
        if should_ocr:
            t0 = time.time()
            try:
                pixmap = page.get_pixmap(dpi=100)
                jpeg_bytes = pixmap.tobytes("jpeg")

                print(f"    OCR page {page_num + 1}/{len(pdf)} of {path.name}...")
                text, extraction_method = _ocr_page(jpeg_bytes)

                if text:
                    ocr_count += 1
                    stats.ocr_pages += 1
                else:
                    # Both local and Gemma OCR returned empty — count as failure.
                    stats.ocr_failed_pages += 1
                    extraction_method = "failed"
                    page_error = "OCR returned empty"
                    stats.errors.append(
                        f"{path.name}:p.{page_num + 1} OCR failed: EmptyResult"
                    )
                    print(f"    ✗ OCR produced no text for page {page_num + 1}")
            except Exception as e:
                stats.ocr_failed_pages += 1
                extraction_method = "failed"
                page_error = type(e).__name__
                stats.errors.append(
                    f"{path.name}:p.{page_num + 1} OCR failed: {type(e).__name__}"
                )
                print(f"    ✗ OCR error on page {page_num + 1}: {type(e).__name__}: {e}")
                text = ""
            page_duration = time.time() - t0
        elif not text:
            # Empty text AND no visual content (or OCR disabled) → blank page.
            # Still count OCR-disabled pages as "blank" from the loader's PoV;
            # the pipeline-wide stat already distinguishes this case via the
            # DISABLE_OCR env var if callers care.
            stats.skipped_blank_pages += 1
            extraction_method = "skipped_blank"

        stats.pages.append(PageLoadResult(
            page=page_num + 1,
            extraction_method=extraction_method,
            text_length=len(text),
            duration_s=page_duration,
            error=page_error,
        ))

        if not text:
            continue

        docs.append(Document(
            text=text,
            metadata={
                "source": path.name,
                "source_path": str(path),
                "page": page_num + 1,
                "total_pages": len(pdf),
                "course": course,
                "category": category,
                "file_type": "pdf",
                "extraction_method": extraction_method,
            }
        ))

    if ocr_count > 0:
        print(f"    ✓ OCR extracted {ocr_count} pages from {path.name}")
    if stats.ocr_failed_pages > 0:
        print(f"    ⚠ OCR failed on {stats.ocr_failed_pages} page(s) of {path.name}")

    pdf.close()

    if return_stats:
        return docs, stats
    return docs


def load_markdown(path: Path, course: str = "", category: str = "") -> List[Document]:
    """Load a markdown file as a single Document."""
    text = path.read_text(encoding="utf-8", errors="replace").strip()

    if not text:
        return []

    return [Document(
        text=text,
        metadata={
            "source": path.name,
            "source_path": str(path),
            "course": course,
            "category": category,
            "file_type": "markdown",
            "extraction_method": "text",
        }
    )]


def load_text(path: Path, course: str = "", category: str = "") -> List[Document]:
    """Load a plain text file (e.g. lecture captions) as a single Document."""
    text = path.read_text(encoding="utf-8", errors="replace").strip()

    if not text:
        return []

    return [Document(
        text=text,
        metadata={
            "source": path.name,
            "source_path": str(path),
            "course": course,
            "category": category,
            "file_type": "text",
            "extraction_method": "text",
        }
    )]


def load_docx(path: Path, course: str = "", category: str = "") -> List[Document]:
    """Load a DOCX file as a single Document."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = DocxDocument(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)

    if not text:
        return []

    return [Document(
        text=text,
        metadata={
            "source": path.name,
            "source_path": str(path),
            "course": course,
            "category": category,
            "file_type": "docx",
            "extraction_method": "text",
        }
    )]


def load_html(path: Path, course: str = "", category: str = "") -> List[Document]:
    """Load an HTML file, extracting text content."""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("beautifulsoup4 is required: pip install beautifulsoup4")

    html_content = path.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(html_content, "html.parser")

    # Remove script and style elements
    for element in soup(["script", "style", "nav", "footer"]):
        element.decompose()

    text = soup.get_text(separator="\n", strip=True)

    if not text:
        return []

    return [Document(
        text=text,
        metadata={
            "source": path.name,
            "source_path": str(path),
            "course": course,
            "category": category,
            "file_type": "html",
            "extraction_method": "text",
        }
    )]


# Extension → loader mapping
LOADERS = {
    ".pdf": load_pdf,
    ".md": load_markdown,
    ".txt": load_text,
    ".docx": load_docx,
    ".html": load_html,
}


def load_file(path: Path, course: str = "", category: str = "") -> List[Document]:
    """Load a file using the appropriate loader based on extension."""
    ext = path.suffix.lower()
    loader = LOADERS.get(ext)

    if loader is None:
        print(f"  ⚠ Unsupported file type: {ext} ({path.name})")
        return []

    try:
        return loader(path, course=course, category=category)
    except Exception as e:
        print(f"  ✗ Error loading {path.name}: {type(e).__name__}: {e}")
        return []


def load_file_with_stats(
    path: Path, course: str = "", category: str = ""
) -> Tuple[List[Document], PdfLoadStats]:
    """
    Load a file and return (docs, stats). Only PDFs produce non-trivial
    stats today; other formats return an empty PdfLoadStats.

    Used by the pipeline so `/api/documents` can surface per-file OCR
    status (ocr_pages / ocr_failed_pages / errors).
    """
    ext = path.suffix.lower()
    loader = LOADERS.get(ext)

    if loader is None:
        print(f"  ⚠ Unsupported file type: {ext} ({path.name})")
        return [], PdfLoadStats()

    try:
        if ext == ".pdf":
            return load_pdf(path, course=course, category=category, return_stats=True)
        return loader(path, course=course, category=category), PdfLoadStats()
    except Exception as e:
        err = f"{path.name}: {type(e).__name__}: {e}"
        print(f"  ✗ Error loading {err}")
        stats = PdfLoadStats()
        stats.errors.append(err)
        return [], stats
